"""
services/document_processor.py
────────────────────────────────
Алгоритмический (без ИИ) парсер документов → структурированный JSON.

Поддерживаемые форматы: DOCX, PDF, PNG/JPG/WEBP (OCR), TXT/MD.

Результат сохраняется в БД (таблица processed_documents) при первом upload,
при повторных запросах — просто читается из кэша. ИИ получает уже готовый
JSON, а не сырой файл, что экономит ~70-90% токенов.

Структура JSON:
{
  "filename": "lecture1.docx",
  "format": "docx",
  "pages": [
    {
      "page": 1,
      "paragraphs": ["текст параграфа 1", ...],
      "images": [
        {"index": 0, "ocr_text": "текст с картинки", "alt": ""}
      ],
      "tables": [
        [["ячейка A1", "ячейка B1"], ["ячейка A2", "ячейка B2"]]
      ]
    }
  ],
  "full_text": "весь текст одной строкой для эмбеддингов"
}
"""

import io
import json
import logging
import re
import zipfile
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


# ═══════════════════════════════════════════════════════════════════════════════
#  PUBLIC ENTRY POINT
# ═══════════════════════════════════════════════════════════════════════════════

def process_document(data: bytes, filename: str) -> dict:
    """
    Parse bytes → structured dict (no AI involved).
    Call once at upload time; cache result in DB.
    """
    ext = Path(filename).suffix.lower()

    if ext == ".docx":
        result = _parse_docx(data, filename)
    elif ext == ".pdf":
        result = _parse_pdf(data, filename)
    elif ext in {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff", ".tif"}:
        result = _parse_image(data, filename)
    elif ext in {".txt", ".md", ".csv"}:
        result = _parse_plaintext(data, filename)
    else:
        raise ValueError(f"Unsupported format: {ext}")

    # Always attach flat full_text for embedding / search
    result["full_text"] = _build_full_text(result)
    return result


def doc_to_prompt_text(doc_json: dict, max_chars: int = 12000) -> str:
    """
    Convert cached JSON → compact text suitable for sending to LLM.
    Much shorter than the raw file: strips XML, deduplicates whitespace,
    respects max_chars budget.
    """
    return doc_json.get("full_text", "")[:max_chars]


# ═══════════════════════════════════════════════════════════════════════════════
#  DOCX PARSER  (python-docx — pure algorithm, zero AI)
# ═══════════════════════════════════════════════════════════════════════════════

def _parse_docx(data: bytes, filename: str) -> dict:
    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError:
        # Fallback: raw XML parse without python-docx
        return _parse_docx_raw_xml(data, filename)

    doc = Document(io.BytesIO(data))
    pages: list[dict] = []
    current_page: dict = {"page": 1, "paragraphs": [], "images": [], "tables": []}

    # ── Paragraphs ────────────────────────────────────────────────────────────
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        current_page["paragraphs"].append(text)

        # Check for page break inside paragraph runs
        for run in para.runs:
            if run._element.xml.find("w:lastRenderedPageBreak") != -1 or \
               run._element.xml.find("w:pageBreak") != -1:
                pages.append(current_page)
                current_page = {
                    "page": len(pages) + 1,
                    "paragraphs": [],
                    "images": [],
                    "tables": [],
                }

    # ── Tables ────────────────────────────────────────────────────────────────
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        if rows:
            current_page["tables"].append(rows)

    # ── Images (inline) ───────────────────────────────────────────────────────
    image_idx = 0
    with zipfile.ZipFile(io.BytesIO(data)) as z:
        image_names = [n for n in z.namelist() if n.startswith("word/media/")]
        for img_name in image_names:
            img_bytes = z.read(img_name)
            ocr_text = _ocr_bytes(img_bytes)
            current_page["images"].append({
                "index": image_idx,
                "ocr_text": ocr_text,
                "alt": Path(img_name).name,
            })
            image_idx += 1

    pages.append(current_page)

    return {"filename": filename, "format": "docx", "pages": pages}


def _parse_docx_raw_xml(data: bytes, filename: str) -> dict:
    """Fallback DOCX parser using only stdlib — extracts text via XML regex."""
    texts = []
    images_ocr = []
    tables = []

    with zipfile.ZipFile(io.BytesIO(data)) as z:
        # Text from word/document.xml
        if "word/document.xml" in z.namelist():
            xml = z.read("word/document.xml").decode("utf-8", errors="ignore")
            found = re.findall(r"<w:t[^>]*>([^<]+)</w:t>", xml)
            texts = [t.strip() for t in found if t.strip()]

        # Images
        img_names = [n for n in z.namelist() if n.startswith("word/media/")]
        for img_name in img_names:
            img_bytes = z.read(img_name)
            ocr = _ocr_bytes(img_bytes)
            if ocr:
                images_ocr.append({"index": len(images_ocr), "ocr_text": ocr, "alt": Path(img_name).name})

    page = {"page": 1, "paragraphs": texts, "images": images_ocr, "tables": tables}
    return {"filename": filename, "format": "docx", "pages": [page]}


# ═══════════════════════════════════════════════════════════════════════════════
#  PDF PARSER
# ═══════════════════════════════════════════════════════════════════════════════

def _parse_pdf(data: bytes, filename: str) -> dict:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages = []

    for page_num, page in enumerate(reader.pages):
        text = (page.extract_text() or "").strip()
        paragraphs = [p.strip() for p in text.split("\n") if p.strip()] if text else []

        # OCR if no text layer
        if not paragraphs:
            ocr = _ocr_pdf_page(page)
            paragraphs = [p for p in ocr.split("\n") if p.strip()]

        # Extract images from page resources
        images = []
        try:
            for img_idx, img_obj in enumerate(page.images):
                ocr_text = _ocr_bytes(img_obj.data)
                images.append({
                    "index": img_idx,
                    "ocr_text": ocr_text,
                    "alt": getattr(img_obj, "name", f"img_{img_idx}"),
                })
        except Exception:
            pass

        pages.append({
            "page": page_num + 1,
            "paragraphs": paragraphs,
            "images": images,
            "tables": [],
        })

    return {"filename": filename, "format": "pdf", "pages": pages}


def _ocr_pdf_page(page) -> str:
    try:
        import pytesseract
        from PIL import Image as PILImage
        pil_image = page.to_image(resolution=150).original
        return pytesseract.image_to_string(pil_image, lang="rus+eng").strip()
    except Exception as e:
        logger.debug("PDF page OCR failed: %s", e)
        return ""


# ═══════════════════════════════════════════════════════════════════════════════
#  IMAGE PARSER
# ═══════════════════════════════════════════════════════════════════════════════

def _parse_image(data: bytes, filename: str) -> dict:
    ocr_text = _ocr_bytes(data)
    page = {
        "page": 1,
        "paragraphs": [p for p in ocr_text.split("\n") if p.strip()],
        "images": [{"index": 0, "ocr_text": ocr_text, "alt": filename}],
        "tables": [],
    }
    return {"filename": filename, "format": "image", "pages": [page]}


# ═══════════════════════════════════════════════════════════════════════════════
#  PLAIN TEXT
# ═══════════════════════════════════════════════════════════════════════════════

def _parse_plaintext(data: bytes, filename: str) -> dict:
    text = data.decode("utf-8", errors="replace").strip()
    paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
    page = {"page": 1, "paragraphs": paragraphs, "images": [], "tables": []}
    return {"filename": filename, "format": "text", "pages": [page]}


# ═══════════════════════════════════════════════════════════════════════════════
#  OCR HELPER
# ═══════════════════════════════════════════════════════════════════════════════

def _ocr_bytes(data: bytes) -> str:
    """Run Tesseract OCR on image bytes. Returns empty string on failure."""
    try:
        import pytesseract
        from PIL import Image as PILImage
        image = PILImage.open(io.BytesIO(data))
        return pytesseract.image_to_string(image, lang="rus+eng").strip()
    except Exception as e:
        logger.debug("OCR failed: %s", e)
        return ""


# ═══════════════════════════════════════════════════════════════════════════════
#  FULL TEXT BUILDER
# ═══════════════════════════════════════════════════════════════════════════════

def _build_full_text(doc: dict) -> str:
    """
    Flatten all pages → single clean string for embeddings and LLM context.
    Order: paragraphs → table cells → image OCR text (per page).
    """
    parts = []
    for page in doc.get("pages", []):
        parts.extend(page.get("paragraphs", []))

        for table in page.get("tables", []):
            for row in table:
                row_text = " | ".join(cell for cell in row if cell)
                if row_text.strip():
                    parts.append(row_text)

        for img in page.get("images", []):
            ocr = img.get("ocr_text", "").strip()
            if ocr:
                parts.append(f"[Изображение: {ocr}]")

    return "\n".join(parts)
